#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
crawl_docx_fulltext.py — 从本地 docx 文件提取南怀瑾著作全文,输出 JSONL。

处理 data/ 目录中的现成文本,补充大全.jpg 中缺失但在网上找不到免费全文的著作。

当前处理:
  - 太湖楞严讲习录 (164章, ~3.5MB docx, 精校本)
  - 太湖版《达摩多罗禅经》音频讲座实录 (~10章, ~3MB docx)
  - 准提法开示系列 (9篇, 分散 docx 文件)

⚠️ 授权闸门:南怀瑾先生著作仍在著作权保护期内。本脚本仅在
   "公版 / 开放许可 / 你已明确获得授权" 的情况下使用。
   运行必须显式加 --i-have-authorization,否则拒绝执行。

断点续传:已输出的文档会被跳过(append 模式)。
"""

import argparse
import json
import os
import re
import sys

from docx import Document

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")
# 如果默认路径不存在,尝试同级 nan/data 目录
if not os.path.isdir(DATA_DIR):
    alt = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "data")
    if os.path.isdir(alt):
        DATA_DIR = alt
OUT_DEFAULT = "docx_documents.jsonl"
ERRORS_DEFAULT = "docx_crawl_errors.jsonl"
TODAY = "2026-07-11"

# ── 预定义清单 ──────────────────────────────────────────
# 格式: {work, category, filename, extract_mode, source_label}
# extract_mode: "lengyan" = 楞严经模式(按"楞严经NNN"分割)
#               "damo" = 达摩禅经模式(按主章节标记分割,跳过TOC)
#               "single" = 单篇模式(全文作为一篇)
PREBUILT = [
    # ── 太湖楞严讲习录 ──
    {
        "work": "太湖楞严讲习录",
        "category": "佛家典籍解读",
        "filename": "⑪精校太湖楞严讲习录1~164完稿（行入南学2024.9.16）.docx",
        "extract_mode": "lengyan",
        "source_label": "行入南学校对组（本地精校本）",
    },
    # ── 太湖版达摩多罗禅经 ──
    {
        "work": "太湖版《达摩多罗禅经》讲座",
        "category": "打坐禅定",
        "filename": "④精校太湖版《达摩多罗禅经》音频讲座实录 行入南学校对组20231222冬至修订版 2.docx",
        "extract_mode": "damo",
        "source_label": "行入南学校对组（本地精校本）",
    },
    # ── 准提法开示系列 ──
    {
        "work": "南师2002年准提法开示录",
        "category": "打坐禅定",
        "filename": "南师2002年准提法开示录（最后校对：2011年10月21日）.docx",
        "extract_mode": "single",
        "source_label": "本地整理稿",
    },
    {
        "work": "南公上师1997年再传准提法于香港（上）",
        "category": "打坐禅定",
        "filename": "南公上師1997年再傳準提法於香港 (上)（20230530修正）.docx",
        "extract_mode": "single",
        "source_label": "胡松年整理稿",
    },
    {
        "work": "南公上师1997年再传准提法于香港（下）",
        "category": "打坐禅定",
        "filename": "南公上師1997年再傳準提法於香港 (下) .docx",
        "extract_mode": "single",
        "source_label": "胡松年整理稿",
    },
    {
        "work": "南师1985年亲领准提法共修开示",
        "category": "打坐禅定",
        "filename": "南懷瑾老師1985年親領「準提法」共修開示（胡松年先生整理）.docx",
        "extract_mode": "single",
        "source_label": "胡松年整理稿",
    },
    {
        "work": "融通白骨禅观与准提法的修持",
        "category": "打坐禅定",
        "filename": "金剛上師南公懷瑾講於1979年底2月底-融通白骨禪觀與準提法的修持.docx",
        "extract_mode": "single",
        "source_label": "胡松年整理稿",
    },
    {
        "work": "殊胜的准提佛母修法",
        "category": "打坐禅定",
        "filename": "殊勝的準提佛母修法.docx",
        "extract_mode": "single",
        "source_label": "本地整理稿",
    },
    {
        "work": "南公准提法已经透露的内密",
        "category": "打坐禅定",
        "filename": "南公「準提法」已經透露的內密 (南師衣中珠之探尋系列).docx",
        "extract_mode": "single",
        "source_label": "本地整理稿",
    },
    {
        "work": "南师开示：宁可执有如须弥山不可落空如芥子许",
        "category": "打坐禅定",
        "filename": "南师开示：宁可执有如须弥山，不可落空如芥子许.docx",
        "extract_mode": "single",
        "source_label": "本地整理稿",
    },
]


def load_done(out_path):
    """读取已有输出,返回已输出的 doc+chapter 集合(断点续传)。"""
    done = set()
    try:
        with open(out_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    key = (obj.get("work", ""), obj.get("chapter_title", ""))
                    done.add(key)
                except json.JSONDecodeError:
                    continue
    except FileNotFoundError:
        pass
    return done


def extract_lengyan(doc):
    """从楞严讲习录 docx 按章节标记拆分。

    章节标记格式: "楞严经001经题", "楞严经002经题及有关资料", ...
    """
    CH_MARKER = re.compile(r'^楞严经(\d{3})(.*)')
    chapters = []  # [(title, paragraph_start_idx, paragraph_end_idx)]

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        m = CH_MARKER.match(t)
        if m:
            ch_num = int(m.group(1))
            ch_desc = m.group(2).strip()
            title = f"楞严经{ch_num:03d}"
            if ch_desc:
                title += f" {ch_desc}"
            chapters.append((title, i, ch_num))

    if not chapters:
        return [("全文", _all_paras_text(doc))]

    # 计算每章结束位置
    results = []
    for ci, (title, start_idx, ch_num) in enumerate(chapters):
        end_idx = chapters[ci + 1][1] if ci + 1 < len(chapters) else len(doc.paragraphs)
        chapter_paras = [doc.paragraphs[j] for j in range(start_idx, end_idx)]
        text = _paras_to_text(chapter_paras)
        results.append((title, text, ch_num))

    return results


def extract_damo(doc):
    """从达摩多罗禅经 docx 按主章节标记拆分。

    章节标记格式: "01《达摩多罗禅经》修行方便道安那般那念退分第一"
    跳过前面的目录部分(第一个标记集合),使用后面的正文部分。

    docx 结构: TOC(para 0~226) + 正文(para 227~end)
    """
    CH_MARKER = re.compile(r'^(\d{2})《达摩多罗禅经》(.*)')
    SEQ_MARKER = re.compile(r'^《达摩多罗禅经》序(\d+)')

    # 找到所有章节标记
    markers = []  # [(title, para_idx)]
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        m = CH_MARKER.match(t)
        if m:
            markers.append((f"{m.group(1)}《达摩多罗禅经》{m.group(2).strip()}", i))
        else:
            m2 = SEQ_MARKER.match(t)
            if m2:
                markers.append((f"《达摩多罗禅经》序{m2.group(1)}", i))

    if not markers:
        return [("全文", _all_paras_text(doc))]

    # 使用后半部分的标记(正文部分)
    # 策略:找到第一个重复的标记,从那里开始
    # TOC 区的标题带 \t页码 后缀(如 "01《达摩多罗禅经》...\t29"),
    # 正文区的标题不带页码。先 strip 页码再比较。
    seen_titles = {}
    content_start = 0
    for mi, (title, idx) in enumerate(markers):
        base = re.sub(r'\t\d+$', '', title)  # 只去掉末尾的制表符+页码
        if base in seen_titles:
            content_start = mi
            break
        seen_titles[base] = mi

    # 使用正文部分的标记
    body_markers = markers[content_start:]
    if not body_markers:
        body_markers = markers

    results = []
    for mi, (title, start_idx) in enumerate(body_markers):
        end_idx = body_markers[mi + 1][1] if mi + 1 < len(body_markers) else len(doc.paragraphs)
        chapter_paras = [doc.paragraphs[j] for j in range(start_idx, end_idx)]
        text = _paras_to_text(chapter_paras)
        results.append((title, text, mi + 1))

    return results


def extract_single(doc):
    """全文作为单篇提取。"""
    text = _all_paras_text(doc)
    return [("全文", text, 1)]


def _paras_to_text(paras):
    """将段落列表转为干净的文本。"""
    lines = []
    for p in paras:
        t = p.text.strip()
        if t:
            lines.append(t)
    text = "\n\n".join(lines)
    # 清理多余空行
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _all_paras_text(doc):
    """获取文档全部文本。"""
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(lines)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def main():
    ap = argparse.ArgumentParser(description="从本地 docx 提取南怀瑾著作全文 → JSONL")
    ap.add_argument("--i-have-authorization", action="store_true",
                    help="确认内容为公版/开放许可/已获授权。必填闸门。")
    ap.add_argument("--out", default=OUT_DEFAULT)
    ap.add_argument("--errors", default=ERRORS_DEFAULT)
    ap.add_argument("--data-dir", default=DATA_DIR,
                    help="docx 文件目录(默认 ../data/)")
    ap.add_argument("--work", default="",
                    help="只处理标题包含该子串的作品(大小写无关)")
    args = ap.parse_args()

    if not args.i_have_authorization:
        sys.exit(
            "本脚本仅用于公版/开放许可/已获授权内容。\n"
            "南怀瑾先生著作仍在著作权保护期内,确认你已获得授权后,加 --i-have-authorization 再运行。"
        )

    done = load_done(args.out)

    works = list(PREBUILT)
    if args.work:
        needle = args.work.lower()
        works = [w for w in works if needle in w["work"].lower()]
        if not works:
            sys.exit(f"没有标题包含 {args.work!r} 的作品。")

    err_fp = open(args.errors, "a", encoding="utf-8")
    out_fp = open(args.out, "a", encoding="utf-8")
    total_new = 0
    total_skip = 0
    total_err = 0
    global_no = 0

    for wi, work_info in enumerate(works, 1):
        work_title = work_info["work"]
        fpath = os.path.join(args.data_dir, work_info["filename"])

        print(f"[{wi}/{len(works)}] {work_title}", flush=True)

        if not os.path.exists(fpath):
            print(f"  !! 文件不存在: {fpath}", flush=True)
            err_fp.write(json.dumps(
                {"stage": "file", "work": work_title, "file": fpath,
                 "error": "文件不存在"}, ensure_ascii=False) + "\n")
            err_fp.flush()
            total_err += 1
            continue

        try:
            doc = Document(fpath)
        except Exception as e:
            print(f"  !! 无法打开: {e}", flush=True)
            err_fp.write(json.dumps(
                {"stage": "open", "work": work_title, "file": fpath,
                 "error": str(e)}, ensure_ascii=False) + "\n")
            err_fp.flush()
            total_err += 1
            continue

        mode = work_info.get("extract_mode", "single")
        try:
            if mode == "lengyan":
                chapters = extract_lengyan(doc)
            elif mode == "damo":
                chapters = extract_damo(doc)
            else:
                chapters = extract_single(doc)
        except Exception as e:
            print(f"  !! 提取失败: {e}", flush=True)
            err_fp.write(json.dumps(
                {"stage": "extract", "work": work_title,
                 "error": str(e)}, ensure_ascii=False) + "\n")
            err_fp.flush()
            total_err += 1
            continue

        print(f"  {len(chapters)} 章/篇", flush=True)

        n_new = 0
        n_skip = 0
        for ch_title, ch_text, ch_num in chapters:
            key = (work_title, ch_title)
            if key in done:
                n_skip += 1
                continue

            global_no += 1
            rec = {
                "work": work_title,
                "category": work_info["category"],
                "work_url": fpath,
                "chapter_no": global_no,
                "chapter_no_in_work": ch_num,
                "chapter_title": ch_title,
                "chapter_url": f"file://{fpath}",
                "char_count": len(ch_text),
                "text": ch_text,
                "retrieved": TODAY,
                "source": work_info["source_label"],
            }
            out_fp.write(json.dumps(rec, ensure_ascii=False) + "\n")
            out_fp.flush()
            done.add(key)
            n_new += 1
            if len(chapters) <= 1:
                print(f"   {ch_title} ({len(ch_text)} 字)", flush=True)
            elif n_new <= 3 or n_new > len(chapters) - 3:
                print(f"   [{ch_num}] {ch_title} ({len(ch_text)} 字)", flush=True)
            elif n_new == 4:
                print(f"   ...", flush=True)

        total_new += n_new
        total_skip += n_skip

    out_fp.close()
    err_fp.close()
    print(f"\n完成:新提取 {total_new},跳过(已存在) {total_skip},失败 {total_err}。")
    print(f"  输出: {args.out}")
    if total_err:
        print(f"  错误: {args.errors}")


if __name__ == "__main__":
    main()
